import datetime
import json
import os
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk

from docx import Document
from openai import OpenAI

# --------------------
# CONFIG
# --------------------
MODEL = "gpt-5-nano"
TEMPLATE_FILE = "Feasibility_Report_Word_Template.docx"
OUTPUT_FILE = "Feasibility_Report_Final.docx"
DATE = datetime.date.today().strftime("%B %d, %Y")

# Keep naming consistent with existing template placeholders.
SECTIONS = [
    "Overview",
    "Business Case",
    "Problem Definition",
    "Value Proposition",
    "Challenges",
    "Success Criteria",
    "Data Availability",
    "Input corpus",
    "Data quality",
    "Technical Feasibility",
    "Analytical complexity",
    "Infrastructure",
    "Integration",
    "Risk Assessment",
    "Mitigation strategies",
    "Recommendations",
    "Recommendation to Proceed",
    "Rationale",
]


# --------------------
# TEMPLATE REPLACER (reused original logic)
# --------------------
def replace_in_paragraph(paragraph, replacements):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)

    replaced = False
    for key, val in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, val)
            replaced = True

    if replaced:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""


def apply_replacements_to_doc(doc, replacements):
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)


class FeasibilityReportApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Feasibility Report Generator")
        self.root.geometry("1200x900")

        self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY", "YOUR_OPENAI_API_KEY"))
        self.generated_data = {section: "" for section in SECTIONS}
        self.approved_sections = set()

        self.section_text_widgets = {}
        self.approve_button_widgets = {}

        self._build_ui()

    def _build_ui(self):
        top_frame = ttk.Frame(self.root, padding=12)
        top_frame.pack(fill="x")

        ttk.Label(top_frame, text="Project Title:").grid(row=0, column=0, sticky="w", padx=(0, 6), pady=4)
        self.project_title_var = tk.StringVar(value="AI Feasibility Study")
        ttk.Entry(top_frame, textvariable=self.project_title_var, width=80).grid(
            row=0, column=1, sticky="ew", pady=4
        )

        ttk.Label(top_frame, text="Author:").grid(row=1, column=0, sticky="w", padx=(0, 6), pady=4)
        self.author_var = tk.StringVar(value="John Doe")
        ttk.Entry(top_frame, textvariable=self.author_var, width=80).grid(
            row=1, column=1, sticky="ew", pady=4
        )

        ttk.Label(top_frame, text="Use Case Description:").grid(
            row=2, column=0, sticky="nw", padx=(0, 6), pady=4
        )
        self.use_case_text = tk.Text(top_frame, width=90, height=8, wrap="word")
        self.use_case_text.insert(
            "1.0",
            "Describe the use case, goals, stakeholders, constraints, and expected outcomes.",
        )
        self.use_case_text.grid(row=2, column=1, sticky="ew", pady=4)

        button_row = ttk.Frame(top_frame)
        button_row.grid(row=3, column=1, sticky="w", pady=(8, 0))

        self.generate_initial_button = ttk.Button(
            button_row,
            text="Generate Initial Report Content",
            command=self.generate_initial_report,
        )
        self.generate_initial_button.pack(side="left")

        self.generate_report_button = ttk.Button(
            button_row,
            text="Generate Report",
            command=self.generate_report_document,
            state="disabled",
        )
        self.generate_report_button.pack(side="left", padx=(10, 0))

        self.status_label = ttk.Label(top_frame, text="Generate initial content to begin.")
        self.status_label.grid(row=4, column=1, sticky="w", pady=(8, 0))

        top_frame.grid_columnconfigure(1, weight=1)

        sections_container = ttk.Frame(self.root, padding=(12, 8, 12, 12))
        sections_container.pack(fill="both", expand=True)

        self.canvas = tk.Canvas(sections_container, highlightthickness=0)
        y_scroll = ttk.Scrollbar(sections_container, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=y_scroll.set)

        y_scroll.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        self.sections_frame = ttk.Frame(self.canvas)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.sections_frame, anchor="nw")

        self.sections_frame.bind("<Configure>", self._on_sections_configure)
        self.canvas.bind("<Configure>", self._on_canvas_configure)

    def _on_sections_configure(self, _event):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _on_canvas_configure(self, event):
        self.canvas.itemconfigure(self.canvas_window, width=event.width)

    def _set_buttons_state(self, enabled):
        state = "normal" if enabled else "disabled"
        self.generate_initial_button.configure(state=state)
        for section in SECTIONS:
            widgets = self.section_text_widgets.get(section)
            if widgets:
                widgets.configure(state=state)
            approve_button = self.approve_button_widgets.get(section)
            if approve_button:
                approve_button.configure(state=state)

    def _build_sections_ui(self):
        for widget in self.sections_frame.winfo_children():
            widget.destroy()

        self.section_text_widgets.clear()
        self.approve_button_widgets.clear()

        for index, section in enumerate(SECTIONS):
            section_frame = ttk.LabelFrame(self.sections_frame, text=section, padding=8)
            section_frame.grid(row=index, column=0, sticky="ew", padx=2, pady=6)
            section_frame.grid_columnconfigure(0, weight=1)

            section_text = tk.Text(section_frame, height=6, wrap="word")
            section_text.insert("1.0", self.generated_data.get(section, ""))
            section_text.grid(row=0, column=0, sticky="ew")

            controls = ttk.Frame(section_frame)
            controls.grid(row=1, column=0, sticky="w", pady=(8, 0))

            improve_button = ttk.Button(
                controls,
                text="Improve",
                command=lambda s=section: self.improve_section(s),
            )
            improve_button.pack(side="left")

            approve_button = ttk.Button(
                controls,
                text="Approve",
                command=lambda s=section: self.approve_section(s),
            )
            approve_button.pack(side="left", padx=(8, 0))

            self.section_text_widgets[section] = section_text
            self.approve_button_widgets[section] = approve_button

    def _validate_inputs(self):
        if not self.project_title_var.get().strip():
            messagebox.showerror("Validation Error", "Project title is required.")
            return False
        if not self.author_var.get().strip():
            messagebox.showerror("Validation Error", "Author is required.")
            return False
        if not self.get_use_case_description():
            messagebox.showerror("Validation Error", "Use Case Description is required.")
            return False
        return True

    def get_use_case_description(self):
        return self.use_case_text.get("1.0", "end").strip()

    def _generate_initial_prompt(self, use_case_description):
        return f"""
Return ONLY valid JSON.

Keys:
{SECTIONS}

Use Case Description:
{use_case_description}
"""

    def _generate_improvement_prompt(self, use_case_description, section_name, current_content, user_feedback):
        return f"""
You are writing a formal government feasibility report. Regenerate ONLY the content for the section \"{section_name}\" based on the following:

Use Case Description:
{use_case_description}

Current Content for {section_name}:
{current_content}

User Feedback for Improvement:
{user_feedback}

Return ONLY the improved text for this section.
"""

    def _call_openai(self, prompt):
        response = self.client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": "You write formal government feasibility reports."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
        )
        return response.choices[0].message.content.strip()

    def _parse_json_response(self, raw_text):
        cleaned = raw_text.strip()

        if cleaned.startswith("```"):
            cleaned = cleaned.strip("`")
            if cleaned.lower().startswith("json"):
                cleaned = cleaned[4:].strip()

        return json.loads(cleaned)

    def generate_initial_report(self):
        if not self._validate_inputs():
            return

        use_case_description = self.get_use_case_description()
        prompt = self._generate_initial_prompt(use_case_description)

        self.status_label.configure(text="Generating initial sections...")
        self.root.update_idletasks()

        try:
            raw = self._call_openai(prompt)
            data = self._parse_json_response(raw)
        except Exception as exc:
            messagebox.showerror("Generation Error", f"Failed to generate report sections.\n\n{exc}")
            self.status_label.configure(text="Generation failed.")
            return

        for section in SECTIONS:
            self.generated_data[section] = str(data.get(section, "")).strip()

        self.approved_sections.clear()
        self._build_sections_ui()
        self._update_approval_status()
        self.status_label.configure(text="Initial content generated. Review each section, then approve.")

    def improve_section(self, section_name):
        if section_name not in self.section_text_widgets:
            return

        feedback = simpledialog.askstring(
            "Improve Section",
            f"How should '{section_name}' be improved?",
            parent=self.root,
        )
        if not feedback:
            return

        use_case_description = self.get_use_case_description()
        current_content = self.section_text_widgets[section_name].get("1.0", "end").strip()

        prompt = self._generate_improvement_prompt(
            use_case_description,
            section_name,
            current_content,
            feedback,
        )

        self.status_label.configure(text=f"Improving: {section_name}...")
        self.root.update_idletasks()

        try:
            improved = self._call_openai(prompt)
        except Exception as exc:
            messagebox.showerror("Improvement Error", f"Failed to improve section.\n\n{exc}")
            self.status_label.configure(text="Section improvement failed.")
            return

        improved = improved.strip()
        self.generated_data[section_name] = improved

        text_widget = self.section_text_widgets[section_name]
        text_widget.delete("1.0", "end")
        text_widget.insert("1.0", improved)

        self.approved_sections.discard(section_name)
        self.approve_button_widgets[section_name].configure(text="Approve")
        self._update_approval_status()
        self.status_label.configure(text=f"Updated {section_name}. Approve it when ready.")

    def approve_section(self, section_name):
        if section_name not in self.section_text_widgets:
            return

        latest_text = self.section_text_widgets[section_name].get("1.0", "end").strip()
        self.generated_data[section_name] = latest_text

        self.approved_sections.add(section_name)
        self.approve_button_widgets[section_name].configure(text="Approved")
        self._update_approval_status()

    def _update_approval_status(self):
        approved_count = len(self.approved_sections)
        total = len(SECTIONS)
        all_approved = approved_count == total
        if all_approved:
            self.status_label.configure(
                text=f"All sections approved ({approved_count}/{total}). Generate Report is enabled."
            )
        else:
            self.status_label.configure(text=f"Sections approved: {approved_count}/{total}")
        self.generate_report_button.configure(state="normal" if all_approved else "disabled")

    def _build_replacements(self):
        replacements = {
            "{{Project Title}}": self.project_title_var.get().strip(),
            "{{Author Name}}": self.author_var.get().strip(),
            "{{Date}}": DATE,
        }

        for section in SECTIONS:
            replacements[f"{{{{{section}}}}}"] = self.generated_data.get(section, "")

        return replacements

    def generate_report_document(self):
        if len(self.approved_sections) != len(SECTIONS):
            messagebox.showerror("Approval Required", "Approve every section before generating the report.")
            return

        output_file = filedialog.asksaveasfilename(
            title="Save Report As",
            defaultextension=".docx",
            initialfile=OUTPUT_FILE,
            filetypes=[("Word Document", "*.docx")],
        )
        if not output_file:
            return

        try:
            doc = Document(TEMPLATE_FILE)
            replacements = self._build_replacements()
            apply_replacements_to_doc(doc, replacements)
            doc.save(output_file)
        except Exception as exc:
            messagebox.showerror("Report Error", f"Failed to generate document.\n\n{exc}")
            return

        messagebox.showinfo("Success", f"Created: {output_file}")


def main():
    root = tk.Tk()
    app = FeasibilityReportApp(root)
    app._set_buttons_state(True)
    root.mainloop()


if __name__ == "__main__":
    main()
